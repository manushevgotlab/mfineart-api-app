#!/usr/bin/env python3
"""Generate MFineArt-Project-Documentation.docx from structured content."""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise SystemExit("Install python-docx: pip install python-docx")

OUTPUT = Path(__file__).parent / "MFineArt-Project-Documentation.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()

    title = doc.add_heading("MFineArt — Project Documentation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Version: 0.0.1-SNAPSHOT  |  Last updated: June 2026  |  Component: REST API backend")
    doc.add_paragraph()

    add_heading(doc, "1. Introduction", 1)
    add_heading(doc, "1.1 Purpose", 2)
    add_para(doc, (
        "MFineArt is the backend for a personal online fine art gallery. It exposes a REST API "
        "consumed by an Angular single-page application serving public visitors and gallery staff "
        "via an /admin area. The API owns business rules, JWT authentication, content lifecycle, "
        "relational data, and AWS S3 image storage."
    ))

    add_heading(doc, "1.2 Technology stack", 2)
    add_table(doc, ["Concern", "Choice"], [
        ["Language", "Java 17"],
        ["Framework", "Spring Boot 3.3.2"],
        ["Persistence", "Spring Data JPA, H2 (dev)"],
        ["Security", "Spring Security, JWT"],
        ["Object storage", "AWS S3 SDK v2"],
        ["Build", "Maven"],
    ])

    add_heading(doc, "2. System architecture", 1)
    add_para(doc, (
        "The Angular frontend communicates over HTTP/JSON with this Spring Boot API. "
        "Public routes call GET endpoints without authentication. The /admin area authenticates "
        "staff with JWT and attaches the Bearer token to mutating requests. The API persists "
        "metadata in a relational database and stores image binaries in S3."
    ))

    add_heading(doc, "3. Domain model", 1)
    add_heading(doc, "3.1 Core entities", 2)
    add_bullets(doc, [
        "ArtCollection — curated group of paintings; single thumbnailUrl; publishable",
        "Painting — artwork with material, dimensions, price, availability; many images",
        "Event — news/exhibitions; many images",
        "Image — S3 URL; belongs to exactly one Painting or Event",
        "AppUser — back-office account (ADMIN or EDITOR)",
    ])

    add_heading(doc, "3.2 Entity hierarchy", 2)
    add_para(doc, (
        "BaseGalleryEntity (id, name, date) → PublishableEntity (contentStatus, publishAt) "
        "→ ArtCollection, Event, Artwork → Painting. Image extends BaseGalleryEntity only."
    ))

    add_heading(doc, "4. Content lifecycle", 1)
    add_para(doc, "Every collection, painting, and event supports four content statuses:")
    add_table(doc, ["Status", "Public?", "Description"], [
        ["DRAFT", "No", "Default on create; staff only"],
        ["SCHEDULED", "When due", "Publishes at publishAt"],
        ["PUBLISHED", "Yes", "Live on public site"],
        ["ARCHIVED", "No", "Retired content"],
    ])

    add_heading(doc, "4.1 Allowed transitions", 2)
    add_bullets(doc, [
        "DRAFT → SCHEDULED, PUBLISHED, ARCHIVED",
        "SCHEDULED → DRAFT, PUBLISHED",
        "PUBLISHED → ARCHIVED",
        "ARCHIVED → DRAFT",
    ])
    add_para(doc, (
        "Status changes use PUT .../content-status endpoints with ContentStatusUpdateDto. "
        "Regular edit endpoints do not change lifecycle status. A scheduled job runs every minute "
        "to promote SCHEDULED items whose publishAt has passed to PUBLISHED."
    ))

    add_heading(doc, "5. Security and authentication", 1)
    add_table(doc, ["Role", "Capabilities"], [
        ["ADMIN", "Full CRUD including DELETE"],
        ["EDITOR", "Create and update (POST, PUT); no DELETE"],
        ["Anonymous", "GET only; sees published content"],
    ])
    add_para(doc, (
        "Login: POST /auth/login returns a JWT. Send Authorization: Bearer <token> on protected "
        "requests. Rules are defined in security/SecurityConfig.java."
    ))

    add_heading(doc, "6. REST API summary", 1)
    add_table(doc, ["Resource", "Base path", "Public GET?", "Staff write?", "Admin DELETE?"], [
        ["Auth", "/auth", "login only", "me", "—"],
        ["Collections", "/collections", "Yes", "Yes", "Yes"],
        ["Paintings", "/paintings", "Yes", "Yes", "Yes"],
        ["Events", "/events", "Yes", "Yes", "Yes"],
        ["Images", "/images", "Yes", "Yes", "Yes"],
    ])

    add_heading(doc, "6.1 Key endpoints", 2)
    add_bullets(doc, [
        "POST /auth/login — obtain JWT",
        "GET /auth/me — current staff user",
        "PUT /collections/content-status — lifecycle",
        "PUT /paintings/painting/content-status — lifecycle",
        "PUT /events/event/content-status — lifecycle",
        "GET /paintings/collection/{collectionId} — paintings in collection",
        "GET /images/painting/{paintingId} — images for painting",
    ])

    add_heading(doc, "7. AWS S3", 1)
    add_para(doc, (
        "Images upload via S3Service. The API stores the public URL in the database. "
        "Collection thumbnails use thumbnailUrl on the collection entity. "
        "S3 object cleanup on delete is not yet implemented."
    ))

    add_heading(doc, "8. Angular integration", 1)
    add_bullets(doc, [
        "Public site: GET without token; render S3 URLs in img tags",
        "/admin: login, store JWT, HTTP interceptor for Authorization header",
        "Route guard on /admin/**; role-based UI for delete actions",
        "Content workflow: contentStatus badges and content-status API calls",
    ])

    add_heading(doc, "9. Configuration", 1)
    add_table(doc, ["Property", "Purpose"], [
        ["jwt.secret", "JWT signing key (min 32 characters)"],
        ["jwt.expiration-ms", "Token lifetime (default 24h)"],
        ["app.cors.allowed-origins", "Angular origin"],
        ["app.content.scheduled-publish-cron", "Auto-publish schedule"],
        ["app.security.default-admin.*", "Bootstrap admin credentials"],
        ["aws.s3.*", "S3 bucket and credentials"],
    ])

    add_heading(doc, "10. Build and run", 1)
    add_bullets(doc, [
        "mvn spring-boot:run — start API on port 8080",
        "mvn clean test — run tests",
        "mvn clean package — build JAR",
        "H2 console: http://localhost:8080/h2-console (dev)",
    ])

    add_heading(doc, "11. Future improvements", 1)
    add_bullets(doc, [
        "Store s3Key; delete S3 objects on entity delete",
        "OpenAPI specification and API versioning",
        "User management endpoints for EDITOR accounts",
        "Image variants (thumbnail, medium, full)",
        "Production database profile (PostgreSQL)",
    ])

    add_heading(doc, "12. Project source reference", 1)
    add_para(doc, "Authoritative security rules: src/main/java/.../security/SecurityConfig.java")
    add_para(doc, "Lifecycle logic: src/main/java/.../service/content/ContentLifecycleService.java")
    add_para(doc, "Database DDL: src/main/resources/dbscripts/create.sql")
    add_para(doc, "Configuration: src/main/resources/application.properties")

    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build()
